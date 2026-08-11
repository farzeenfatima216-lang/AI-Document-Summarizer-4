import io
import re
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency
    PdfReader = None

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None

try:
    import fitz
except Exception:  # pragma: no cover - optional dependency
    fitz = None

try:
    import numpy as np
except Exception:  # pragma: no cover - optional dependency
    np = None

try:
    import easyocr
    OCR_READER = easyocr.Reader(["en", "ur"], gpu=False, verbose=False)
except Exception:  # pragma: no cover - optional dependency
    easyocr = None
    OCR_READER = None

try:
    from PIL import Image, ImageOps, ImageFilter
except Exception:  # pragma: no cover - optional dependency
    Image = None
    ImageOps = None
    ImageFilter = None


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".json"}


class OCRUnavailableError(RuntimeError):
    """Raised when a scanned PDF needs OCR but its optional tools are missing."""


def normalize_filename(name: Optional[str]) -> str:
    if not name:
        return "uploaded_file"
    return Path(name).name


def _prepare_ocr_image(image):
    if image.mode != "L":
        image = image.convert("L")
    if ImageOps is not None:
        image = ImageOps.autocontrast(image)
        image = ImageOps.equalize(image)
    if ImageFilter is not None:
        image = image.filter(ImageFilter.MedianFilter(size=3))
        image = image.filter(ImageFilter.SHARPEN)
    if hasattr(Image, "Resampling"):
        resample = Image.Resampling.LANCZOS
    elif hasattr(Image, "ANTIALIAS"):
        resample = Image.ANTIALIAS
    elif hasattr(Image, "LANCZOS"):
        resample = Image.LANCZOS
    else:
        resample = 1
    image = image.resize((int(image.width * 2), int(image.height * 2)), resample)
    return image


def _is_scanned_page_text(text: str) -> bool:
    lower = text.lower()
    no_whitespace = re.sub(r"\s+", "", lower)
    alpha_numeric = re.findall(r"[a-z0-9]", lower)
    alpha_num_ratio = len(alpha_numeric) / max(1, len(text))
    word_count = len(re.findall(r"\w+", lower))

    if "camscanner" in lower or "scanned by" in lower or "scan" in lower and "pdf" in lower:
        return True

    if len(no_whitespace) < 50:
        return True

    if word_count < 15 and alpha_num_ratio < 0.4:
        return True

    if len(text.strip().splitlines()) <= 2 and len(no_whitespace) < 120:
        return True

    if re.search(r"[\x00-\x1f\x7f-\x9f]", text):
        return True

    return False


def _clean_scanned_text(text: str) -> str:
    if not text:
        return text
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        lower = stripped.lower()
        if re.fullmatch(r"(?i)(camscanner|scanned by|scan|scanned|pdf|copy|page \d+ of \d+)", lower):
            continue
        if re.search(r"(?i)(camscanner|scanned by|scan|scanned|pdf|copy)", stripped) and len(re.sub(r"\s+", "", stripped)) < 30:
            continue
        cleaned_line = re.sub(r"(?i)\b(camscanner|scanned by|scan|scanned|pdf|copy)\b", "", stripped)
        cleaned_line = re.sub(r"[\s\n]{2,}", " ", cleaned_line).strip()
        if not cleaned_line:
            continue
        if len(re.sub(r"\s+", "", cleaned_line)) < 4:
            continue
        lines.append(cleaned_line)
    cleaned = "\n".join(lines).strip()
    cleaned = re.sub(r"[\n\s]{2,}", "\n", cleaned)
    return cleaned


def _is_good_ocr_text(text: str) -> bool:
    cleaned = _clean_scanned_text(text)
    return bool(cleaned and len(re.sub(r"\s+", "", cleaned)) >= 8)


def _threshold_image(image, threshold=150):
    if Image is None:
        return image
    return image.point(lambda p: 0 if p < threshold else 255, mode="1").convert("L")


def _extract_text_from_image(image):
    image = _prepare_ocr_image(image)
    reader = OCR_READER
    if reader is None or np is None:
        return ""
    try:
        raw_texts = reader.readtext(np.array(image), detail=0, paragraph=True)
        ocr_text = "\n".join(raw_texts).strip()
        print("OCR RESULT >>>", ocr_text[:500])
        return _clean_scanned_text(ocr_text)
    except Exception:
        return ""


def _extract_pdf_text_with_fitz(file_bytes: bytes) -> List[str]:
    if fitz is None:
        raise RuntimeError("PDF text extraction fallback requires PyMuPDF. Install it with pip install PyMuPDF")

    document = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        return [((page.get_text() or "").strip()) for page in document]
    finally:
        document.close()


def _extract_pdf_pages(file_bytes: bytes) -> Tuple[List[str], bool, List[int]]:
    if PdfReader is None and fitz is None:
        raise RuntimeError(
            "PDF support requires pypdf or PyMuPDF. Install at least one of these packages."
        )

    pages: List[str] = []
    if PdfReader is not None:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = [(page.extract_text() or "").strip() for page in reader.pages]
        except Exception as exc:
            if fitz is None:
                raise RuntimeError(
                    "PDF support requires pypdf or PyMuPDF. pypdf failed: " + str(exc)
                ) from exc
            pages = _extract_pdf_text_with_fitz(file_bytes)
    else:
        pages = _extract_pdf_text_with_fitz(file_bytes)

    weak_pages: List[int] = []
    for index, text in enumerate(pages):
        cleaned = _clean_scanned_text(text)
        if _is_scanned_page_text(text) or len(cleaned.strip()) < 80 or "camscanner" in text.lower():
            weak_pages.append(index)

    if not weak_pages:
        return pages, False, []

    if fitz is None or Image is None or OCR_READER is None:
        raise OCRUnavailableError(
            "This PDF appears to be scanned. OCR requires PyMuPDF, Pillow, and easyocr (pip install PyMuPDF Pillow easyocr)."
        )

    ocr_pages: List[int] = []
    document = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        for page_index in weak_pages:
            page = None
            if hasattr(document, "load_page"):
                try:
                    page = document.load_page(page_index)
                except Exception:
                    page = None
            if page is None:
                try:
                    page = document[page_index]
                except Exception:
                    try:
                        page = list(document)[page_index]
                    except Exception:
                        page = None

            if page is None:
                pages[page_index] = ""
                continue

            page_number = page_index + 1
            best_text = ""

            for zoom in (3, 4, 5):
                try:
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                    image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                except Exception:
                    continue

                ocr_text = _extract_text_from_image(image)
                print(f"Page {page_number}, zoom {zoom}, OCR length = {len(ocr_text)}")
                print(ocr_text[:300])

                if _is_good_ocr_text(ocr_text) and len(ocr_text) > len(best_text):
                    best_text = ocr_text

                if len(best_text) > 50:
                    break

            if not best_text:
                raw_text = ""
                if hasattr(page, "get_text"):
                    try:
                        raw_text = _clean_scanned_text(page.get_text().strip())
                    except Exception:
                        raw_text = ""

                if _is_good_ocr_text(raw_text):
                    best_text = raw_text

            pages[page_index] = best_text
            if best_text:
                ocr_pages.append(page_number)
    finally:
        try:
            document.close()
        except Exception:
            pass

    return pages, True, ocr_pages


def extract_text_from_file(
    file_name: Optional[str], file_bytes: bytes, pdf_pages: Optional[List[str]] = None
) -> str:
    name = normalize_filename(file_name).lower()

    if name.endswith((".txt", ".md", ".json")):
        return file_bytes.decode("utf-8", errors="ignore").strip()

    if name.endswith(".pdf"):
        pages = pdf_pages if pdf_pages is not None else _extract_pdf_pages(file_bytes)[0]
        return "\n".join(text for text in pages if text).strip()

    if name.endswith(".docx"):
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
                candidate_names = [
                    "word/document.xml",
                    "word/footnotes.xml",
                    "word/endnotes.xml",
                ]
                for candidate in candidate_names:
                    if candidate in archive.namelist():
                        xml_bytes = archive.read(candidate)
                        text = re.sub(r"<[^>]+>", " ", xml_bytes.decode("utf-8", errors="ignore"))
                        text = re.sub(r"\s+", " ", text).strip()
                        if text:
                            return text
        except Exception:
            pass

        if Document is not None:
            try:
                document = Document(io.BytesIO(file_bytes))
                paragraphs = [
                    para.text.strip()
                    for para in document.paragraphs
                    if para.text and para.text.strip()
                ]
                return "\n".join(paragraphs).strip()
            except Exception as exc:
                raise RuntimeError(f"DOCX parsing failed: {exc}") from exc

        raise RuntimeError("DOCX parsing failed because the file is not a valid Word document")

    raise ValueError("Unsupported file type. Use txt, md, json, pdf, or docx.")


def _section_name(text: str, current: str) -> str:
    line = text.strip().strip(":")
    lower = line.lower()
    if lower in {"camscanner", "scanned", "scan", "pdf", "copy"}:
        return current
    if line and len(line) <= 80 and (
        text.strip().endswith(":")
        or line.isupper()
        or (not current and len(line.split()) <= 8)
    ):
        return line
    return current


def extract_document_units(
    file_name: Optional[str], file_bytes: bytes, pdf_pages: Optional[List[str]] = None
) -> List[Dict]:
    """Extract text units while retaining page and paragraph locations."""
    name = normalize_filename(file_name).lower()
    units = []

    if name.endswith(".pdf"):
        pages = pdf_pages if pdf_pages is not None else _extract_pdf_pages(file_bytes)[0]
        for page_number, page_text in enumerate(pages, start=1):
            section = ""
            paragraphs = [part.strip() for part in re.split(r"\n\s*\n", page_text) if part.strip()]
            if not paragraphs and page_text.strip():
                paragraphs = [page_text.strip()]
            for paragraph_number, paragraph in enumerate(paragraphs, start=1):
                section = _section_name(paragraph, section)
                units.append({
                    "text": paragraph,
                    "page_number": page_number,
                    "paragraph_number": paragraph_number,
                    "section": section or "Document",
                })
        return units

    if name.endswith(".docx") and Document is not None:
        document = Document(io.BytesIO(file_bytes))
        section = ""
        paragraph_number = 0
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            paragraph_number += 1
            section = _section_name(text, section)
            units.append({
                "text": text,
                "page_number": None,
                "paragraph_number": paragraph_number,
                "section": section or "Document",
            })
        if units:
            return units

    raw_text = file_bytes.decode("utf-8", errors="ignore").strip()
    section = ""
    for paragraph_number, paragraph in enumerate(
        (part.strip() for part in re.split(r"\n\s*\n", raw_text)), start=1
    ):
        if not paragraph:
            continue
        section = _section_name(paragraph, section)
        units.append({
            "text": paragraph,
            "page_number": None,
            "paragraph_number": paragraph_number,
            "section": section or "Document",
        })
    return units


def build_chunk_records(units: List[Dict], chunk_size: int = 1000) -> List[Dict]:
    records = []
    for unit in units:
        words = unit["text"].split()
        current = []
        current_length = 0
        for word in words:
            if current and current_length + len(word) + 1 > chunk_size:
                records.append({**unit, "text": " ".join(current)})
                current = [word]
                current_length = len(word)
            else:
                current.append(word)
                current_length += len(word) + 1
        if current:
            records.append({**unit, "text": " ".join(current)})
    return records

def split_text_into_chunks(text: str, chunk_size: int = 700, chunk_overlap: int = 100) -> List[str]:
    if not text or not text.strip():
        return []

    words = re.split(r"\s+", text.strip())
    if not words:
        return []

    chunks = []
    current = []
    current_len = 0

    for word in words:
        if current and current_len + len(word) + 1 > chunk_size:
            chunks.append(" ".join(current))
            current = [word]
            current_len = len(word)
        else:
            current.append(word)
            current_len += len(word) + 1

    if current:
        chunks.append(" ".join(current))

    if chunk_overlap > 0 and len(chunks) > 1:
        overlapped = []
        for index, chunk in enumerate(chunks):
            if index == 0:
                overlapped.append(chunk)
            else:
                previous = chunks[index - 1]
                overlap_words = previous.split()[-chunk_overlap // 5:]
                overlapped.append(" ".join(overlap_words + chunk.split()))
        return overlapped

    return chunks


def prepare_document_payload(file_name: Optional[str], file_bytes: bytes) -> dict:
    pdf_pages = None
    used_ocr = False
    ocr_pages = []
    scanned_pdf = False
    ocr_warning = None

    if normalize_filename(file_name).lower().endswith(".pdf"):
        try:
            pdf_pages, used_ocr, ocr_pages = _extract_pdf_pages(file_bytes)
            scanned_pdf = used_ocr
        except OCRUnavailableError as exc:
            scanned_pdf = True
            ocr_warning = str(exc)
            if PdfReader is not None:
                reader = PdfReader(io.BytesIO(file_bytes))
                pdf_pages = [(page.extract_text() or "").strip() for page in reader.pages]
            else:
                pdf_pages = []

    text = extract_text_from_file(file_name, file_bytes, pdf_pages=pdf_pages)
    units = extract_document_units(file_name, file_bytes, pdf_pages=pdf_pages)
    chunk_records = build_chunk_records(units)
    chunks = [record["text"] for record in chunk_records]
    extraction_method = "ocr" if used_ocr else ("scanned_pdf" if scanned_pdf else "embedded_text")

    payload = {
        "text": text,
        "chunks": chunks,
        "chunk_records": chunk_records,
        "chunk_count": len(chunks),
        "text_length": len(text),
        "extraction_method": extraction_method,
        "scanned_pdf": scanned_pdf,
        "ocr_pages": ocr_pages,
    }
    if ocr_warning:
        payload["ocr_warning"] = ocr_warning
    return payload
