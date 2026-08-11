import io
import importlib
import json
import os
import re
import sys
import traceback
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from threading import Thread
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from tempfile import SpooledTemporaryFile
from time import perf_counter
from urllib.parse import parse_qs, urlparse

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

import utils
from rag import RAGUnavailableError, retrieve_relevant_chunks
from prompt_engineering import build_prompt

ROOT = Path(__file__).resolve().parent
PORT = int(os.environ.get("PORT", "8000"))

try:
    import cgi  # type: ignore
except Exception:  # pragma: no cover - Python 3.13+
    cgi = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency
    PdfReader = None

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None

try:
    from fpdf import FPDF
except Exception:  # pragma: no cover - optional dependency
    FPDF = None

ChatGroq = None
LLM = None
GROQ_MODEL = os.environ.get("GROQ_MODEL", "llama-3.1-8b-instant")

if os.getenv("GROQ_API_KEY"):
    try:
        from langchain_groq import ChatGroq
        LLM = ChatGroq(model=GROQ_MODEL, temperature=0)
        print(f"[LLM] enabled, model={GROQ_MODEL}")
    except Exception as exc:  # pragma: no cover - optional dependency / runtime safety
        print(f"[LLM] DISABLED - could not initialise ChatGroq: {exc!r}")
        ChatGroq = None
        LLM = None
else:
    print("[LLM] DISABLED - GROQ_API_KEY is not set. "
          "Summaries and answers will use the non-AI fallback.")


def split_into_chunks(text, chunk_size=1000):
    if not text or not text.strip():
        return []
    words = text.split()
    chunks = []
    current = []
    current_len = 0
    for word in words:
        if current_len + len(word) + 1 > chunk_size and current:
            chunks.append(" ".join(current))
            current = [word]
            current_len = len(word)
        else:
            current.append(word)
            current_len += len(word) + 1
    if current:
        chunks.append(" ".join(current))
    return chunks


def retrieve_context(question, context, chunks=None, top_k=3):
    candidate_chunks = chunks
    if candidate_chunks is None:
        candidate_chunks = split_into_chunks(context, chunk_size=800)
    try:
        start = perf_counter()
        retrieved = retrieve_relevant_chunks(question, candidate_chunks, top_k=top_k)
        elapsed = perf_counter() - start
        retrieved_docs = [item.get("file_name") or item.get("document_name") or "Unknown" for item in retrieved]
        print(f"[RAG] retrieval_time={elapsed:.3f}s top_k={top_k} chunks={len(candidate_chunks)} retrieved_docs={retrieved_docs}")
        return "\n\n".join(item["chunk"] for item in retrieved), retrieved
    except (RAGUnavailableError, ValueError) as exc:
        print(f"[RAG] retrieval fallback: {exc}")
        return context, []


def _normalize_file_items(file_item):
    if file_item is None:
        return []
    if isinstance(file_item, list):
        return file_item
    return [file_item]


def build_citations(retrieved):
    return [
        {
            "document_name": item.get("document_name") or item.get("file_name") or item.get("source") or "Unknown",
            "file_name": item.get("file_name") or item.get("document_name") or item.get("source") or "Unknown",
            "page_number": item.get("page_number"),
            "paragraph_number": item.get("paragraph_number"),
            "section": item.get("section") or "Document",
            "score": item.get("score", 0),
            "excerpt": item.get("chunk", ""),
        }
        for item in retrieved
    ]


def fallback_sources(chunks, question=""):
    sources = chunks or []
    question_tokens = set(
        token for token in re.findall(r"[a-z0-9]+", question.lower())
        if len(token) > 1
    )
    records = []
    for item in sources:
        chunk = item.get("text", "") if isinstance(item, dict) else item
        chunk_tokens = set(re.findall(r"[a-z0-9]+", chunk.lower()))
        score = (
            len(question_tokens & chunk_tokens) / len(question_tokens)
            if question_tokens
            else 0.0
        )
        records.append({
            "chunk": chunk,
            "document_name": item.get("document_name") if isinstance(item, dict) else None,
            "file_name": item.get("file_name") if isinstance(item, dict) else None,
            "page_number": item.get("page_number") if isinstance(item, dict) else None,
            "paragraph_number": (
                item.get("paragraph_number") if isinstance(item, dict) else None
            ),
            "section": item.get("section", "Document") if isinstance(item, dict) else "Document",
            "score": score,
        })
    records.sort(key=lambda item: item["score"], reverse=True)
    return records[:3]


def _format_summary_bullets(sentences, count=4):
    bullets = []
    for sentence in sentences:
        if sentence:
            bullets.append(f"- {sentence}")
            if len(bullets) >= count:
                break
    return bullets


def _summarize_with_variant(text, variant=None):
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    sentences = [s.strip() for s in sentences if s.strip()]
    if not sentences:
        return "No readable text was found in the uploaded file."

    if variant == "concise":
        return " ".join(sentences[:2]) if len(sentences) >= 2 else sentences[0]

    if variant == "audience_first":
        summary = " ".join(sentences[:3])
        return f"This document is for a busy reader. {summary}"

    if variant == "structured_grounded":
        bullets = _format_summary_bullets(sentences, count=4)
        return "\n".join(bullets) if bullets else sentences[0]

    return " ".join(sentences[:3])


def _invoke_llm_with_timeout(prompt, timeout=15):
    if not LLM:
        return None

    executor = getattr(_invoke_llm_with_timeout, "executor", None)
    if executor is None:
        executor = ThreadPoolExecutor(max_workers=1)
        _invoke_llm_with_timeout.executor = executor

    try:
        future = executor.submit(LLM.invoke, prompt)
        response = future.result(timeout=timeout)
        return getattr(response, "content", str(response)).strip()
    except FutureTimeoutError:
        print(f"[LLM] TIMEOUT after {timeout}s - falling back to non-AI output")
        try:
            future.cancel()
        except Exception:
            pass
        return None
    except Exception as exc:
        print(f"[LLM] CALL FAILED: {type(exc).__name__}: {exc}")
        try:
            future.cancel()
        except Exception:
            pass
        return None


def build_summary(text, variant=None):
    if LLM:
        try:
            prompt = build_prompt("summary", variant_name=variant, document=text)
            summary = _invoke_llm_with_timeout(prompt, timeout=15)
            if summary:
                return summary
        except Exception as exc:
            print(f"[LLM] build_summary failed: {type(exc).__name__}: {exc}")

    return _summarize_with_variant(text, variant=variant)


def _extract_key_points_with_variant(text, variant=None):
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    sentences = [s.strip() for s in sentences if s.strip()]
    if not sentences:
        return ["No readable content found in the uploaded file."]

    if variant == "bullet_points":
        bullets = [s for s in sentences[:5]]
        return [f"- {point}" for point in bullets]

    if variant == "executive_summary":
        summary_lines = sentences[:5]
        return [f"{idx + 1}. {line}" for idx, line in enumerate(summary_lines)]

    if variant == "evidence_tagged":
        points = []
        for sentence in sentences[:5]:
            tail = " (from document)"
            points.append(f"{sentence[:150].rstrip('.')}.{tail}")
        return points

    return [s[:180] + ("..." if len(s) > 180 else "") for s in sentences[:5]]


def build_key_points(text, variant=None):
    if LLM:
        try:
            prompt = build_prompt("key_points", variant_name=variant, document=text)
            response = LLM.invoke(prompt)
            content = getattr(response, "content", str(response)).strip()
            if content:
                return [line.strip(" -•") for line in content.splitlines() if line.strip()]
        except Exception as exc:
            print(f"[LLM] build_key_points failed: {type(exc).__name__}: {exc}")

    return _extract_key_points_with_variant(text, variant=variant)


def build_answer(question, context, chunks=None, variant=None):
    if not question or not context:
        return "Please provide both a question and document content."


    start_total = perf_counter()
    context, retrieved = retrieve_context(question, context, chunks, top_k=3)
    retrieval_time = perf_counter() - start_total
    print(f"Retrieved Documents: {[item.get('file_name') or item.get('document_name') or 'Unknown' for item in retrieved]}")
    for item in retrieved:
        print("Retrieved Document Name", item.get("file_name") or item.get("document_name") or "Unknown")
        print("Retrieved Page Number", item.get("page_number"))
        print("Retrieved Chunk Text", (item.get("chunk") or item.get("text", ""))[:100])

    if retrieved:
        candidate_records = retrieved
    elif chunks is not None:
        candidate_records = [
            item if isinstance(item, dict) else {"chunk": item}
            for item in chunks
            if item and (item.get("text", "") if isinstance(item, dict) else item)
        ]
    else:
        candidate_records = []

    candidate_chunks = []
    for item in candidate_records:
        if isinstance(item, dict):
            chunk_text = item.get("chunk") or item.get("text") or ""
        else:
            chunk_text = item or ""
        if chunk_text:
            candidate_chunks.append(chunk_text)
    relevant_context = "\n\n".join(candidate_chunks)
    if not relevant_context:
        relevant_context = context[:4000]

    if LLM:
        try:
            prompt = build_prompt(
                "qa",
                variant_name=variant,
                question=question,
                context=relevant_context[:8000],
            )
            start_llm = perf_counter()
            response = _invoke_llm_with_timeout(prompt, timeout=15)
            llm_time = perf_counter() - start_llm
            if response:
                total_time = perf_counter() - start_total
                print(
                    f"[QA] total_time={total_time:.3f}s retrieval_time={retrieval_time:.3f}s llm_time={llm_time:.3f}s"
                )
                print(f"[LLM] response_time={llm_time:.3f}s prompt_tokens~{len(prompt)//4}")
                return response
        except Exception as exc:
            print(f"[LLM] invoke failed: {exc}")

    if candidate_chunks:
        keyword_tokens = [
            token for token in re.findall(r"[a-z0-9]+", question.lower())
            if token not in {
                "what", "which", "when", "where", "who", "why", "how",
                "does", "do", "did", "are", "is", "the", "this",
                "that", "can", "could", "would", "should", "please",
                "tell", "me",
            }
            and len(token) > 1
        ]
        best_match = ""
        best_score = 0
        for chunk in candidate_chunks:
            if not isinstance(chunk, str) or not chunk.strip():
                continue
            lower_chunk = chunk.lower()
            # Count whole-word matches only. Plain str.count() would match "ai"
            # inside "said" or "rain", which matters now that two-letter
            # acronyms (AI, ML, UK) are allowed as search terms.
            score = sum(
                len(re.findall(rf"\b{re.escape(token)}\b", lower_chunk))
                for token in keyword_tokens
            )
            if score > best_score or (
                score > 0 and score == best_score and len(chunk) > len(best_match)
            ):
                best_score = score
                best_match = chunk
        if best_match:
            sentences = [
                sentence.strip()
                for sentence in re.split(r"(?<=[.!?])\s+", best_match)
                if sentence.strip()
            ]
            if sentences:
                best_sentence = max(
                    sentences,
                    key=lambda sentence: sum(
                        sentence.lower().count(token) for token in keyword_tokens
                    ),
                )
            else:
                best_sentence = best_match[:700].strip()

            source_doc = None
            source_page = None
            source_para = None
            source_section = None
            if candidate_records and isinstance(candidate_records[0], dict):
                source_doc = candidate_records[0].get("file_name")
                source_page = candidate_records[0].get("page_number")
                source_para = candidate_records[0].get("paragraph_number")
                source_section = candidate_records[0].get("section")

            citation = []
            if source_doc:
                citation.append(f"Document: {source_doc}")
            if source_page is not None:
                citation.append(f"Page {source_page}")
            if source_para is not None:
                citation.append(f"Para {source_para}")
            if source_section:
                citation.append(f"Section: {source_section}")
            citation_text = " | ".join(citation)

            total_time = perf_counter() - start_total
            print(
                f"[QA] total_time={total_time:.3f}s retrieval_time={retrieval_time:.3f}s heuristic_fallback"
            )
            if citation_text:
                return f"{best_sentence[:700].strip()}\n\nSource: {citation_text}"
            return best_sentence[:700].strip()

    total_time = perf_counter() - start_total
    print(
        f"[QA] total_time={total_time:.3f}s retrieval_time={retrieval_time:.3f}s no_answer"
    )
    return "Not found in the document."

    best_match = ""
    best_score = 0
    for chunk in candidate_chunks:
        if not isinstance(chunk, str) or not chunk.strip():
            continue
        lower_chunk = chunk.lower()
        score = sum(lower_chunk.count(token) for token in keyword_tokens)
        if score > best_score or (
            score > 0 and score == best_score and len(chunk) > len(best_match)
        ):
            best_score = score
            best_match = chunk

    if best_match:
        sentences = [
            sentence.strip()
            for sentence in re.split(r"(?<=[.!?])\s+", best_match)
            if sentence.strip()
        ]
        if sentences:
            best_sentence = max(
                sentences,
                key=lambda sentence: sum(
                    sentence.lower().count(token) for token in keyword_tokens
                ),
            )
            return best_sentence[:700].strip()
        return best_match[:700].strip()

    return "Not found in the document."


def parse_form_data(body_bytes, content_type):
    if not body_bytes:
        return {}

    if "multipart/form-data" not in content_type:
        return {}

    boundary = None
    for part in content_type.split(";"):
        if part.strip().startswith("boundary="):
            boundary = part.split("=", 1)[1].strip().strip('"')
            break

    if not boundary:
        raise ValueError("Multipart boundary is missing from the request")

    boundary_bytes = boundary.encode("utf-8")
    raw_parts = body_bytes.split(b"--" + boundary_bytes)
    result = {}

    for raw_part in raw_parts:
        if not raw_part or raw_part in (b"--", b"--\r\n", b"\r\n"):
            continue

        chunk = raw_part.strip(b"\r\n")
        if not chunk:
            continue

        if b"\r\n\r\n" in chunk:
            header_bytes, data = chunk.split(b"\r\n\r\n", 1)
        elif b"\n\n" in chunk:
            header_bytes, data = chunk.split(b"\n\n", 1)
        else:
            continue

        headers = {}
        for line in header_bytes.splitlines():
            if b":" in line:
                key, value = line.split(b":", 1)
                headers[key.decode("latin-1").strip().lower()] = value.decode("latin-1").strip()

        disposition = headers.get("content-disposition", "")
        name = None
        filename = None
        for item in disposition.split(";"):
            item = item.strip()
            if item.startswith("name="):
                name = item.split("=", 1)[1].strip().strip('"')
            elif item.startswith("filename="):
                filename = item.split("=", 1)[1].strip().strip('"')

        if name is None:
            continue

        data_bytes = data.rstrip(b"\r\n")
        value = type("FormValue", (), {})()
        value.name = name
        value.filename = filename
        if filename:
            value.file = SpooledTemporaryFile()
            value.file.write(data_bytes)
            value.file.seek(0)
        else:
            value.value = data_bytes.decode("utf-8", errors="ignore")

        if name in result:
            existing = result[name]
            if isinstance(existing, list):
                existing.append(value)
            else:
                result[name] = [existing, value]
        else:
            result[name] = value

    return result


# Simple in-memory job store for async analysis
from uuid import uuid4
JOBS = {}


def process_document_job(job_id, file_items, prompt_variants):
    """Background worker to process uploaded files and update JOBS[job_id] with progress and result."""
    try:
        JOBS[job_id]["status"] = "processing"
        JOBS[job_id]["progress"] = 5
        documents = []
        combined_texts = []
        combined_records = []
        file_names = []

        for idx, file_item in enumerate(file_items):
            JOBS[job_id]["progress"] = min(40, 10 + idx * 5)
            if not getattr(file_item, "file", None):
                continue
            file_name = os.path.basename(file_item.filename or f"uploaded_{idx}")
            file_bytes = file_item.file.read()
            if isinstance(file_bytes, str):
                file_bytes = file_bytes.encode("utf-8")
            try:
                document = utils.prepare_document_payload(file_name, file_bytes)
            except Exception as e:
                # Fallback: don't fail the whole job for one bad file. Record error and continue.
                JOBS[job_id].setdefault("file_errors", []).append({"file": file_name, "error": str(e)})
                document = {"text": "", "chunk_records": [], "extraction_method": "failed", "scanned_pdf": False, "ocr_pages": []}
            documents.append(document)
            file_names.append(file_name)
            combined_texts.append(f"DOCUMENT: {file_name}\n{document['text']}")
            for record in document["chunk_records"]:
                record_with_source = dict(record)
                record_with_source["file_name"] = file_name
                record_with_source["document_name"] = file_name
                record_with_source["section"] = record_with_source.get("section", "Document")
                combined_records.append(record_with_source)

        JOBS[job_id]["progress"] = 55
        text = "\n\n".join(combined_texts).strip()
        chunks = [record["text"] for record in combined_records if record.get("text")]
        chunk_records = combined_records

        print(f"[JOB] job_id={job_id} uploaded_files={len(file_items)} processed_files={len(file_names)} total_text_length={len(text)} total_chunks={len(chunks)}")
        for idx, document in enumerate(documents):
            print(f"[JOB] job_id={job_id} doc_index={idx + 1} file_name={file_names[idx]} text_length={len(document['text'])} chunk_count={len(document['chunk_records'])}")

        JOBS[job_id]["progress"] = 70
        summary = build_summary(text, variant=prompt_variants.get("summary"))
        JOBS[job_id]["progress"] = 85
        key_points = build_key_points(text, variant=prompt_variants.get("key_points"))

        trimmed_records = [
            {
                "text": r.get("text", "")[:400],
                "page_number": r.get("page_number"),
                "paragraph_number": r.get("paragraph_number"),
                "section": r.get("section", "Document"),
                "file_name": r.get("file_name", ""),
                "document_name": r.get("document_name", "")
            }
            for r in chunk_records[:30]
        ]

        payload = {
            "status": "success",
            "message": "File processed successfully",
            "file_names": file_names,
            "document_count": len(file_names),
            "text": text[:8000],
            "preview": text[:3000],
            "summary": summary,
            "key_points": key_points,
            "chunks": chunks[:50],
            "chunk_records": trimmed_records,
            "chunk_count": len(chunks),
            "text_length": len(text),
            "extraction_method": "multi_document" if len(file_names) > 1 else (documents[0]["extraction_method"] if documents else "Text"),
            "scanned_pdf": any(document.get("scanned_pdf") for document in documents),
            "ocr_pages": [page for document in documents for page in document.get("ocr_pages", [])],
            "prompt_variants": prompt_variants,
        }

        JOBS[job_id]["result"] = payload
        JOBS[job_id]["progress"] = 100
        JOBS[job_id]["status"] = "completed"
    except Exception as exc:  # noqa: BLE001
        JOBS[job_id]["status"] = "error"
        JOBS[job_id]["error"] = str(exc)
        JOBS[job_id]["trace"] = traceback.format_exc()



def parse_urlencoded_form(body_bytes):
    if not body_bytes:
        return {}
    return {key: values[0] if len(values) == 1 else values for key, values in parse_qs(body_bytes.decode("utf-8")).items()}


def extract_text(file_name, file_bytes):
    return extract_text_from_file(file_name, file_bytes)


class Handler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        self.end_headers()

    def end_headers(self):
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        try:
            super().end_headers()
        except (ConnectionAbortedError, BrokenPipeError, OSError) as e:
            # Client disconnected while sending headers; log and stop trying to write.
            print(f"[WARN] end_headers failed: {e}")
            try:
                # Best-effort: flush and ignore
                self.wfile.flush()
            except Exception:
                pass

    def do_GET(self):
        parsed = urlparse(self.path)

        if parsed.path == "/api/health":
            self.send_json(200, {
                "status": "ok",
                "message": "Backend is running"
            })
            return

        if parsed.path == "/favicon.ico":
            self.send_response(204)
            self.send_header("Content-Type", "image/x-icon")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
            self.send_header("Access-Control-Allow-Headers", "Content-Type")
            self.end_headers()
            return
        if parsed.path == "/api/analyze":
            self.send_json(405, {"error": "Use POST to upload a file."})
            return
        if parsed.path == "/api/analyze-status":
            qs = parse_qs(parsed.query)
            job_id = qs.get("job", [None])[0]
            if not job_id or job_id not in JOBS:
                self.send_json(404, {"error": "Job not found"})
                return
            job = JOBS[job_id]
            resp = {
                "job": job_id,
                "status": job.get("status", "unknown"),
                "progress": job.get("progress", 0),
            }
            if job.get("status") == "completed":
                resp["result"] = job.get("result")
            if job.get("status") == "error":
                resp["error"] = job.get("error")
                resp["trace"] = job.get("trace")
            self.send_json(200, resp)
            return
        super().do_GET()

    def do_POST(self):
        parsed = urlparse(self.path)
        if parsed.path == "/api/analyze":
            self.handle_analyze()
            return
        if parsed.path == "/api/ask":
            self.handle_ask()
            return
        if parsed.path == "/api/export":
            self.handle_export()
            return
        self.send_error(404)

    def handle_analyze(self):
        try:
            content_length = int(self.headers.get("Content-Length", "0"))
            body_bytes = self.rfile.read(content_length)
            content_type = self.headers.get("Content-Type", "").lower()

            prompt_variants = {
                "summary": None,
                "key_points": None,
            }
            file_items = []

            def _get_form_value(field):
                if field is None:
                    return None
                if isinstance(field, list):
                    field = field[0]
                return getattr(field, "value", field)

            if "multipart/form-data" in content_type:
                form = parse_form_data(body_bytes, content_type)
                file_items = _normalize_file_items(form.get("file"))
                prompt_variants["summary"] = (
                    _get_form_value(form.get("summary_prompt_variant"))
                    or _get_form_value(form.get("summary_variant"))
                )
                prompt_variants["key_points"] = (
                    _get_form_value(form.get("key_points_prompt_variant"))
                    or _get_form_value(form.get("key_points_variant"))
                )
            elif content_type.startswith("application/x-www-form-urlencoded"):
                form = parse_urlencoded_form(body_bytes)
                prompt_variants["summary"] = form.get("summary_prompt_variant") or form.get("summary_variant")
                prompt_variants["key_points"] = form.get("key_points_prompt_variant") or form.get("key_points_variant")
            else:
                self.send_json(400, {"error": "No file uploaded. Please choose a PDF, DOCX, or text file first."})
                return

            if not file_items:
                self.send_json(400, {"error": "No file uploaded. Please choose a PDF, DOCX, or text file first."})
                return

            print(f"[ANALYZE] Received analyze request with {len(file_items)} file(s)")
            documents = []
            combined_texts = []
            combined_chunks = []
            combined_records = []
            file_names = []
            print("STEP 1 - Reload utils")
            importlib.reload(utils)
            for file_item in file_items:
                if not getattr(file_item, "file", None):
                    continue
                file_name = os.path.basename(file_item.filename or "uploaded_file")
                file_bytes = file_item.file.read()
                if isinstance(file_bytes, str):
                    file_bytes = file_bytes.encode("utf-8")
                try:
                    document = utils.prepare_document_payload(file_name, file_bytes)
                except ValueError as exc:
                    self.send_json(400, {"error": str(exc)})
                    return
                except utils.OCRUnavailableError as exc:
                    document = None
                    error_message = str(exc)
                    self.send_json(422, {"error": error_message})
                    return
                except Exception as exc:
                    error_message = str(exc)
                    if "PDF support requires pypdf" in error_message:
                        self.send_json(422, {"error": "PDF extraction requires pypdf. Install it with pip install pypdf."})
                        return
                    self.send_json(500, {"error": f"Document extraction failed: {error_message}", "trace": traceback.format_exc()})
                    return
                documents.append(document)
                file_names.append(file_name)
                combined_texts.append(f"DOCUMENT: {file_name}\n{document['text']}")
                for record in document["chunk_records"]:
                    record_with_source = dict(record)
                    record_with_source["file_name"] = file_name
                    record_with_source["document_name"] = file_name
                    record_with_source["section"] = record_with_source.get("section", "Document")
                    combined_records.append(record_with_source)
                    if record_with_source.get("text"):
                        combined_chunks.append(record_with_source["text"])
            text = "\n\n".join(combined_texts).strip()
            chunks = combined_chunks
            chunk_records = combined_records

            print("Uploaded Files:", len(file_items))
            print("Processed Documents:", len(documents))
            print("Combined Text Length:", len(text))
            print("Total Chunks:", len(chunks))
            for idx, document in enumerate(documents):
                print(file_names[idx])
                print("text_length", len(document["text"]))
                print("chunk_count", len(document["chunk_records"]))

            print("STEP 3 - Starting summary generation")
            summary = build_summary(text, variant=prompt_variants["summary"])
            print("STEP 3 - Summary generated")

            print("STEP 4 - Starting key points generation")
            key_points = build_key_points(text, variant=prompt_variants["key_points"])
            print("STEP 4 - Key points generated")

            # Trim chunk_records to avoid oversized response (ConnectionAbortedError)
            trimmed_records = [
                {
                    "text": r.get("text", "")[:400],
                    "page_number": r.get("page_number"),
                    "paragraph_number": r.get("paragraph_number"),
                    "section": r.get("section", "Document"),
                    "file_name": r.get("file_name", ""),
                }
                for r in chunk_records[:30]
            ]
            payload = {
                "status": "success",
                "message": "File processed successfully",
                "file_names": file_names,
                "document_count": len(file_names),
                "text": text,
                "preview": text[:3000],
                "summary": summary,
                "key_points": key_points,
                "chunks": chunks[:50],
                "chunk_records": chunk_records,
                "chunk_records_preview": trimmed_records,
                "chunk_count": len(chunks),
                "text_length": len(text),
                "extraction_method": "multi_document" if len(file_names) > 1 else document["extraction_method"],
                "scanned_pdf": any(document.get("scanned_pdf") for document in documents),
                "ocr_pages": [page for document in documents for page in document.get("ocr_pages", [])],
                "prompt_variants": prompt_variants,
            }
            print(f"[ANALYZE] Completed analyze response: file_names={file_names}, chunk_count={len(chunks)}, text_length={len(text)}")
            self.send_json(200, payload)
        except Exception as exc:  # noqa: BLE001
            error_trace = traceback.format_exc()
            print(f"[ANALYZE ERROR] {exc}\n{error_trace}")
            self.send_json(500, {"error": f"Upload failed: {exc}", "trace": error_trace})

    def handle_ask(self):
        try:
            content_type = self.headers.get("Content-Type", "")
            content_length = int(self.headers.get("Content-Length", "0"))
            body_bytes = self.rfile.read(content_length)

            qa_variant = None
            if content_type.startswith("application/json"):
                raw_body = body_bytes.decode("utf-8")
                payload = json.loads(raw_body) if raw_body else {}
                question = payload.get("question", "")
                context = payload.get("context", "") or payload.get("text", "")
                chunks = payload.get("chunks", [])
                chunks = payload.get("chunk_records", chunks)
                qa_variant = payload.get("qa_variant")
            else:
                if cgi is not None:
                    form = cgi.FieldStorage(
                        fp=io.BytesIO(body_bytes),
                        headers=self.headers,
                        environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": content_type},
                    )
                    question = form.getvalue("question", "")
                    context = form.getvalue("context", "")
                    chunks = form.getvalue("chunks", [])
                else:
                    form = parse_urlencoded_form(body_bytes)
                    question = form.get("question", "")
                    context = form.get("context", "")
                    chunks = form.get("chunks", [])

            answer = build_answer(question, context, chunks, variant=qa_variant)
            _, retrieved = retrieve_context(question, context, chunks)
            displayed_retrieved = retrieved or fallback_sources(chunks, question)
            self.send_json(
                200,
                {
                    "status": "success",
                    "answer": answer,
                    "question": question,
                    "retrieval_method": (
                        "FAISS RAG" if retrieved else "keyword fallback"
                    ),
                    "retrieved_chunks": displayed_retrieved,
                    "citations": build_citations(displayed_retrieved),
                },
            )
        except Exception as exc:  # noqa: BLE001
            self.send_json(500, {"error": str(exc)})

    def handle_export(self):
        try:
            content_type = self.headers.get("Content-Type", "")
            content_length = int(self.headers.get("Content-Length", "0"))
            body_bytes = self.rfile.read(content_length)

            if not content_type.startswith("application/json"):
                self.send_json(400, {"error": "Export requires JSON request body."})
                return

            payload = json.loads(body_bytes.decode("utf-8") if body_bytes else "{}")
            summary = payload.get("summary", "")
            key_points = payload.get("key_points", [])
            text = payload.get("text", "")
            retrieved_chunks = payload.get("retrieved_chunks", [])
            file_name = payload.get("file_name", "document-summary")
            file_format = payload.get("format", "txt").lower()

            if file_format not in {"txt", "docx", "pdf"}:
                self.send_json(400, {"error": "Invalid export format. Use txt, docx, or pdf."})
                return

            placeholder_texts = {
                "Your summary will appear here after processing.",
                "No summary available.",
                "No summary available",
            }
            if (not summary or summary in placeholder_texts) and text:
                summary = build_summary(text)
            if (not key_points or key_points == [] or all(not point.strip() for point in key_points)) and text:
                key_points = build_key_points(text)

            if not summary and text:
                summary = build_summary(text)

            if not summary and not text:
                self.send_json(400, {"error": "No summary or text available for export."})
                return

            file_bytes, content_type = self.build_export_file(summary, key_points, retrieved_chunks, file_name, file_format)
            actual_format = file_format
            if content_type.startswith("text/plain") and file_format != "txt":
                print(f"[EXPORT] {file_format} unavailable (missing library) - delivering .txt instead")
                actual_format = "txt"
            download_name = f"{Path(file_name).stem}.{actual_format}"

            self.send_response(200)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Length", str(len(file_bytes)))
            self.send_header("Content-Disposition", f"attachment; filename=\"{download_name}\"")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
            self.send_header("Access-Control-Allow-Headers", "Content-Type")
            self.end_headers()
            self.wfile.write(file_bytes)
        except Exception as exc:  # noqa: BLE001
            self.send_json(500, {"error": str(exc)})

    def build_export_file(self, summary, key_points, retrieved_chunks, file_name, file_format):
        file_name = Path(file_name).stem or "document-summary"
        if file_format == "txt":
            text = self.render_txt_export(summary, key_points, retrieved_chunks)
            return text.encode("utf-8"), "text/plain;charset=utf-8"

        if file_format == "docx":
            if Document is None:
                text = self.render_txt_export(summary, key_points, retrieved_chunks)
                return text.encode("utf-8"), "text/plain;charset=utf-8"
            return self.render_docx_export(summary, key_points, retrieved_chunks, file_name)

        if file_format == "pdf":
            if FPDF is None:
                text = self.render_txt_export(summary, key_points, retrieved_chunks)
                return text.encode("utf-8"), "text/plain;charset=utf-8"
            return self.render_pdf_export(summary, key_points, retrieved_chunks, file_name)

        raise ValueError("Unsupported export format")

    def render_txt_export(self, summary, key_points, retrieved_chunks):
        lines = ["Summary:", summary.strip(), "", "Key points:"]
        if isinstance(key_points, list):
            lines.extend([f"- {point}" for point in key_points if point])
        else:
            lines.append(str(key_points))

        if retrieved_chunks:
            lines.extend(["", "Retrieved context:"])
            for chunk in retrieved_chunks:
                chunk_text = chunk.get("chunk") if isinstance(chunk, dict) else str(chunk)
                chunk_source = chunk.get("file_name") if isinstance(chunk, dict) else None
                chunk_section = chunk.get("section") if isinstance(chunk, dict) else None
                if chunk_source or chunk_section:
                    source_line = " – ".join(filter(None, [chunk_source, chunk_section]))
                    lines.append(f"Source: {source_line}")
                if chunk_text:
                    lines.append(chunk_text.strip())
                    lines.append("")
        return "\n".join(lines).strip()

    def render_docx_export(self, summary, key_points, retrieved_chunks, file_name):
        document = Document()
        document.add_heading(file_name, level=1)
        document.add_heading("Summary", level=2)
        for paragraph in summary.split("\n"):
            document.add_paragraph(paragraph)
        document.add_heading("Key points", level=2)
        if isinstance(key_points, list):
            for point in key_points:
                document.add_paragraph(point, style="List Bullet")
        else:
            document.add_paragraph(str(key_points))

        if retrieved_chunks:
            document.add_heading("Retrieved context", level=2)
            for chunk in retrieved_chunks:
                chunk_text = chunk.get("chunk") if isinstance(chunk, dict) else str(chunk)
                chunk_source = chunk.get("file_name") if isinstance(chunk, dict) else None
                chunk_section = chunk.get("section") if isinstance(chunk, dict) else None
                if chunk_source or chunk_section:
                    source_line = " – ".join(filter(None, [chunk_source, chunk_section]))
                    document.add_paragraph(source_line)
                document.add_paragraph(chunk_text)
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def render_pdf_export(self, summary, key_points, retrieved_chunks, file_name):
        def safe(text):
            return text.encode("latin-1", errors="replace").decode("latin-1")

        pdf = FPDF()
        pdf.set_auto_page_break(True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, safe(file_name), ln=1)
        pdf.ln(5)
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Summary", ln=1)
        pdf.set_font("Arial", "", 11)
        for line in summary.split("\n"):
            pdf.multi_cell(0, 7, safe(line))
        pdf.ln(5)
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, "Key points", ln=1)
        pdf.set_font("Arial", "", 11)
        if isinstance(key_points, list):
            for point in key_points:
                pdf.multi_cell(0, 7, safe(f"- {point}"))
        else:
            pdf.multi_cell(0, 7, safe(str(key_points)))
        if retrieved_chunks:
            pdf.ln(5)
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, "Retrieved context", ln=1)
            pdf.set_font("Arial", "", 10)
            for chunk in retrieved_chunks:
                chunk_text = chunk.get("chunk") if isinstance(chunk, dict) else str(chunk)
                chunk_source = chunk.get("file_name") if isinstance(chunk, dict) else None
                chunk_section = chunk.get("section") if isinstance(chunk, dict) else None
                if chunk_source or chunk_section:
                    source_line = " – ".join(filter(None, [chunk_source, chunk_section]))
                    pdf.multi_cell(0, 7, safe(source_line))
                pdf.multi_cell(0, 7, safe(chunk_text))
                pdf.ln(2)
        pdf_bytes = pdf.output(dest="S")
        if isinstance(pdf_bytes, str):
            pdf_bytes = pdf_bytes.encode("latin-1")
        return pdf_bytes, "application/pdf"

    def send_json(self, status_code, payload):
        try:
            body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        except Exception as exc:  # noqa: BLE001
            body = json.dumps({"error": str(exc)}).encode("utf-8")

        try:
            self.send_response(status_code)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
            self.send_header("Access-Control-Allow-Headers", "Content-Type")
            self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
            self.send_header("Pragma", "no-cache")
            self.send_header("Expires", "0")
            try:
                self.end_headers()
                self.wfile.write(body)
                self.wfile.flush()
            except (ConnectionAbortedError, BrokenPipeError, OSError) as e:
                # Client disconnected during write; log and ignore.
                print(f"[WARN] client disconnected while sending JSON: {e}")
            except Exception as e:  # noqa: BLE001
                print(f"[ERROR] send_json unexpected error: {e}")
        except Exception as e:  # noqa: BLE001
            print(f"[ERROR] send_json failed before writing: {e}")


if __name__ == "__main__":
    host = os.environ.get("HOST", "0.0.0.0")
    with ThreadingHTTPServer((host, PORT), Handler) as httpd:
        print(f"Serving frontend and backend at http://{host}:{PORT}/")
        httpd.serve_forever()
