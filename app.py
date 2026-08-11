"""Streamlit interface for the AI Document Summarizer.

This is the alternative UI. The main UI is the custom HTML frontend served by
`serve_frontend.py`. Both share the same extraction, prompting, and retrieval code.

Run with:  streamlit run app.py
"""

import io
import os
from pathlib import Path

import streamlit as st

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None

try:
    from fpdf import FPDF
except Exception:  # pragma: no cover - optional dependency
    FPDF = None

try:
    from langchain_groq import ChatGroq
except Exception:  # pragma: no cover - optional dependency
    ChatGroq = None

from serve_frontend import build_summary, build_key_points
from utils import prepare_document_payload
from rag import RAGUnavailableError, retrieve_relevant_chunks
from prompt_engineering import build_prompt, get_prompt_variant_names


# ---------------------------------------------------------------------------
# Export helpers
#
# These MUST be defined before they are called. In the previous version they
# were defined at the bottom of the file but called near the top, which raised
# NameError and crashed the app on every run.
# ---------------------------------------------------------------------------

def create_txt_export(summary, key_points):
    export_lines = ["Summary:", (summary or "").strip(), "", "Key points:"]
    if isinstance(key_points, list):
        export_lines.extend([f"- {point}" for point in key_points if point])
    else:
        export_lines.append(str(key_points))
    return "\n".join(export_lines).encode("utf-8")


def create_docx_export(summary, key_points, title):
    if Document is None:
        return None
    doc = Document()
    doc.add_heading(title or "Document summary", level=1)
    doc.add_heading("Summary", level=2)
    for paragraph in (summary or "").split("\n"):
        doc.add_paragraph(paragraph)
    doc.add_heading("Key points", level=2)
    if isinstance(key_points, list):
        for point in key_points:
            doc.add_paragraph(str(point), style="List Bullet")
    else:
        doc.add_paragraph(str(key_points))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def create_pdf_export(summary, key_points, title):
    if FPDF is None:
        return None

    def safe(text):
        # FPDF's core fonts are latin-1 only; drop characters it cannot encode.
        return str(text).encode("latin-1", "replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, safe(title or "Document summary"), ln=1)
    pdf.ln(5)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Summary", ln=1)
    pdf.set_font("Arial", "", 11)
    for line in (summary or "").split("\n"):
        pdf.multi_cell(0, 7, safe(line))
    pdf.ln(5)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Key points", ln=1)
    pdf.set_font("Arial", "", 11)
    if isinstance(key_points, list):
        for point in key_points:
            pdf.multi_cell(0, 7, safe(f"- {point}"))
    else:
        pdf.multi_cell(0, 7, safe(key_points))
    pdf_bytes = pdf.output(dest="S")
    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode("latin-1")
    return pdf_bytes


# ---------------------------------------------------------------------------
# LLM setup
#
# The API key is read from the environment. It is never hardcoded here.
# Set it with:  export GROQ_API_KEY=...   (or in your host's secrets panel)
# ---------------------------------------------------------------------------

GROQ_MODEL = os.environ.get("GROQ_MODEL", "llama-3.1-8b-instant")


@st.cache_resource(show_spinner=False)
def get_llm():
    if not os.getenv("GROQ_API_KEY"):
        return None, "GROQ_API_KEY is not set - using the non-AI fallback."
    if ChatGroq is None:
        return None, "langchain-groq is not installed - using the non-AI fallback."
    try:
        return ChatGroq(model=GROQ_MODEL, temperature=0), None
    except Exception as exc:
        return None, f"Could not start the Groq client ({exc}) - using the non-AI fallback."


llm, llm_warning = get_llm()


# ---------------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------------

st.set_page_config(page_title="AI Document Summarizer", page_icon="📄", layout="wide")

st.title("AI Document Summarizer and Q&A System")
st.write(
    "Upload a PDF, DOCX, or text file to generate a summary, extract key points, "
    "and ask questions based on the document."
)

if llm_warning:
    st.warning(llm_warning)
else:
    st.caption(f"AI enabled · model: {GROQ_MODEL}")

with st.sidebar:
    st.header("Prompt settings")
    summary_variant = st.selectbox("Summary style", get_prompt_variant_names("summary"))
    key_points_variant = st.selectbox("Key points style", get_prompt_variant_names("key_points"))
    qa_variant = st.selectbox("Q&A style", get_prompt_variant_names("qa"))

uploaded_file = st.file_uploader(
    "Upload PDF, DOCX, or text file",
    type=["pdf", "docx", "txt", "md", "json"],
)

if uploaded_file is not None:
    # Only reprocess when a different file is uploaded.
    signature = f"{uploaded_file.name}:{uploaded_file.size}"
    if st.session_state.get("signature") != signature:
        with st.spinner("Extracting text..."):
            payload = prepare_document_payload(uploaded_file.name, uploaded_file.getvalue())

        text = payload["text"]
        st.session_state["signature"] = signature
        st.session_state["text"] = text
        st.session_state["chunks"] = payload["chunks"]
        st.session_state["chunk_records"] = payload.get("chunk_records", payload["chunks"])
        st.session_state["file_name"] = Path(uploaded_file.name).stem

        with st.spinner("Generating summary..."):
            if llm:
                try:
                    prompt = build_prompt("summary", variant_name=summary_variant, document=text[:8000])
                    response = llm.invoke(prompt)
                    st.session_state["summary"] = getattr(response, "content", str(response)).strip()
                except Exception as exc:
                    st.error(f"Summary generation failed: {exc}")
                    st.session_state["summary"] = build_summary(text, variant=summary_variant)
            else:
                st.session_state["summary"] = build_summary(text, variant=summary_variant)

        with st.spinner("Extracting key points..."):
            if llm:
                try:
                    prompt = build_prompt("key_points", variant_name=key_points_variant, document=text[:8000])
                    response = llm.invoke(prompt)
                    content = getattr(response, "content", str(response))
                    st.session_state["key_points"] = [
                        line.strip(" -•") for line in content.splitlines() if line.strip()
                    ]
                except Exception as exc:
                    st.error(f"Key point extraction failed: {exc}")
                    st.session_state["key_points"] = build_key_points(text, variant=key_points_variant)
            else:
                st.session_state["key_points"] = build_key_points(text, variant=key_points_variant)

text = st.session_state.get("text", "")
chunks = st.session_state.get("chunks", [])
chunk_records = st.session_state.get("chunk_records", [])
summary_text = st.session_state.get("summary", "")
key_points = st.session_state.get("key_points", [])

if text:
    with st.expander(f"Document text ({len(text):,} characters)"):
        st.write(text or "No readable text could be extracted from the uploaded file.")

    with st.expander(f"Document chunks ({len(chunks)})"):
        for index, chunk in enumerate(chunks):
            st.markdown(f"**Chunk {index + 1}**")
            st.write(chunk)
            st.divider()

if summary_text:
    st.subheader("Document Summary")
    st.write(summary_text)

    st.subheader("Key Points")
    for point in key_points:
        st.write(f"- {point}")

    # -----------------------------------------------------------------------
    # Export
    # -----------------------------------------------------------------------
    st.subheader("Export")
    format_type = st.selectbox("Export format", ["TXT", "DOCX", "PDF"])
    file_stem = st.session_state.get("file_name", "summary")

    if format_type == "TXT":
        export_bytes = create_txt_export(summary_text, key_points)
        mime_type = "text/plain"
    elif format_type == "DOCX":
        export_bytes = create_docx_export(summary_text, key_points, file_stem)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        export_bytes = create_pdf_export(summary_text, key_points, file_stem)
        mime_type = "application/pdf"

    if export_bytes is not None:
        st.download_button(
            label=f"Export {format_type}",
            data=export_bytes,
            file_name=f"{file_stem}.{format_type.lower()}",
            mime=mime_type,
        )
    else:
        missing = "python-docx" if format_type == "DOCX" else "fpdf"
        st.warning(f"{format_type} export unavailable because `{missing}` is not installed.")

    # -----------------------------------------------------------------------
    # Q&A
    #
    # This block used to be indented under an `else:`, so it only ran when
    # there was NO summary - meaning it never ran at all. It now sits at the
    # top level and reads its data from session state.
    # -----------------------------------------------------------------------
    st.subheader("Ask a Question")
    query = st.text_input("Enter your question")

    if query:
        try:
            retrieved = retrieve_relevant_chunks(query, chunk_records, top_k=3)
            context = "\n\n".join(item["chunk"] for item in retrieved)
            st.caption(f"Retrieved {len(retrieved)} relevant chunks with FAISS")
        except (RAGUnavailableError, ValueError) as exc:
            retrieved = []
            context = "\n".join(str(c) for c in chunks[:3])
            st.warning(str(exc))

        if retrieved:
            with st.expander("Sources"):
                for item in retrieved:
                    page = item.get("page_number") or "N/A"
                    paragraph = item.get("paragraph_number") or "N/A"
                    section = item.get("section") or "Document"
                    st.info(
                        f"Page {page} | Paragraph {paragraph} | Section: {section}\n\n"
                        f"{item['chunk']}"
                    )

        st.subheader("Answer")
        if llm:
            with st.spinner("Finding the best answer..."):
                try:
                    final_prompt = build_prompt(
                        "qa", variant_name=qa_variant, context=context, question=query
                    )
                    response = llm.invoke(final_prompt)
                    st.write(getattr(response, "content", str(response)))
                except Exception as exc:
                    st.error(f"Answer generation failed: {exc}")
        else:
            from serve_frontend import build_answer

            st.write(build_answer(query, context, chunk_records, variant=qa_variant))

elif uploaded_file is None:
    st.info("Upload a document to get started.")
